import fitz  # PyMuPDF
import uuid
from pathlib import Path
from typing import List, Dict
import pandas as pd
import docx

from pipeline.chunker import RecursiveChunker
from pipeline.embedder import Embedder
from storage.chroma_store import ChromaStore
from storage.file_store import register_document
from core.config import settings
import logging

from pipeline.ocr import ocr_page
from pipeline.table_extractor import extract_tables_from_page
from pipeline.version_manager import handle_version_on_upload
from pipeline.bm25_index import rebuild_index

logger = logging.getLogger(__name__)

def extract_pdf_pages(pdf_path: Path) -> List[Dict]:
    pages = []
    doc = fitz.open(str(pdf_path))
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if len(text) < 50:
            logger.info(f"Page {page_num+1} has sparse text ({len(text)} chars), attempting OCR...")
            text = ocr_page(page)
        table_md = extract_tables_from_page(pdf_path, page_num + 1)
        if table_md:
            text = text + "\n\n[TABLES ON THIS PAGE]\n" + table_md
        pages.append({"page_number": page_num + 1, "text": text})
    doc.close()
    return pages

def extract_text_pages(file_path: Path) -> List[Dict]:
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1", errors="replace")
    return [{"page_number": 1, "text": text}]

def extract_docx_pages(file_path: Path) -> List[Dict]:
    doc = docx.Document(str(file_path))
    text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join([cell.text for cell in row.cells])
    return [{"page_number": 1, "text": text}]

def extract_csv_xlsx_pages(file_path: Path, is_csv=True) -> List[Dict]:
    if is_csv:
        df = pd.read_csv(str(file_path))
    else:
        df = pd.read_excel(str(file_path))
    text = df.to_markdown(index=False)
    return [{"page_number": 1, "text": text}]

def ingest_document(file_bytes: bytes, original_filename: str, user_id: str = "") -> Dict:
    """
    Full ingestion pipeline:
    1. Save Document
    2. Extract pages (routes by ext)
    3. Chunk
    4. Embed
    5. Store in ChromaDB
    6. Register in manifest
    Returns document info dict.
    """
    archived_id = handle_version_on_upload(original_filename)
    if archived_id:
        logger.info(f"Archived previous version of {original_filename} (id={archived_id})")

    doc_id = str(uuid.uuid4())
    ext = Path(original_filename).suffix.lower()
    if not ext:
        ext = ".txt" # fallback
        
    file_path = settings.UPLOAD_DIR / f"{doc_id}{ext}"
    file_path.write_bytes(file_bytes)
    
    file_size_kb = len(file_bytes) / 1024
    logger.info(f"Ingesting {original_filename} ({file_size_kb:.1f} KB)")
    
    if ext == ".pdf":
        pages = extract_pdf_pages(file_path)
    elif ext in [".md", ".txt"]:
        pages = extract_text_pages(file_path)
    elif ext in [".docx", ".doc"]:
        pages = extract_docx_pages(file_path)
    elif ext == ".csv":
        pages = extract_csv_xlsx_pages(file_path, is_csv=True)
    elif ext in [".xlsx", ".xls", ".sheet"]:
        pages = extract_csv_xlsx_pages(file_path, is_csv=False)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
        
    total_pages = len(pages)
    
    chunker = RecursiveChunker(chunk_size=settings.CHUNK_SIZE, chunk_overlap=settings.CHUNK_OVERLAP)
    chunks = chunker.chunk_pages(pages)
    
    if not chunks:
        raise ValueError(f"No text could be extracted from {original_filename}")
    
    texts = [c["text"] for c in chunks]
    embeddings = Embedder.embed(texts)
    
    ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [
        {
            "doc_id": doc_id,
            "doc_name": original_filename,
            "page_number": c["page_number"],
            "chunk_index": c["chunk_index"]
        }
        for c in chunks
    ]
    if user_id:
        for m in metadatas:
            m["user_id"] = user_id
            
    ChromaStore.add_chunks(ids=ids, embeddings=embeddings, documents=texts, metadatas=metadatas)
    rebuild_index()
    
    register_document(
        doc_id=doc_id,
        doc_name=original_filename,
        total_pages=total_pages,
        total_chunks=len(chunks),
        file_size_kb=file_size_kb,
        user_id=user_id
    )
    
    logger.info(f"Ingested {len(chunks)} chunks from {total_pages} pages")
    
    return {
        "doc_id": doc_id,
        "doc_name": original_filename,
        "total_pages": total_pages,
        "total_chunks": len(chunks)
    }

extract_pages = extract_pdf_pages
